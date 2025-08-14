"""
Datasheet ingestion utilities for the Simplex RAG engine.

This module provides functions to extract component definitions from
PDF, Word and Excel documents.  It uses open‑source libraries such as
PyMuPDF (``fitz``) and python‑docx to parse text and tables.  If
these libraries are unavailable, it falls back to plain text extraction
and basic regex heuristics.  For complex tables or ambiguous cases,
the ingestion process can delegate extraction to a language model via
batch API calls through the ``LLMInterface``.

The ingestion workflow operates in batches to avoid rate limits when
invoking external LLMs.  It writes a JSONL file of extraction tasks
for the OpenAI batch API and later reads the results back to enrich
the database.
"""

from __future__ import annotations

import json
import logging
import os
import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    import pdfplumber
except ImportError:
    pdfplumber = None  # type: ignore
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore

try:
    import docx
except ImportError:
    docx = None  # type: ignore

try:
    import pandas as pd
except ImportError:
    pd = None  # type: ignore

from ..data_models import Component, ComponentType, Protocol, CertificationType
from ..config import settings
from ..llm_interface import LLMInterface

logger = logging.getLogger(__name__)


def chunk_text(text: str, max_chars: int = 2000, overlap: int = 200) -> List[str]:
    """
    Break a long text into smaller chunks for LLM processing.

    Splits on paragraph boundaries (double newlines) and ensures that
    each chunk does not exceed ``max_chars``.  Adjacent chunks have
    ``overlap`` characters of overlap to preserve context.  If the
    paragraph splitting yields segments larger than ``max_chars``, the
    segment is further split at word boundaries.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        # If adding this paragraph exceeds max_chars, flush current chunk
        if len(current) + len(para) + 2 > max_chars:
            if current:
                chunks.append(current.strip())
                # Start new chunk with overlap from end of current
                overlap_text = current[-overlap:] if overlap > 0 else ""
                current = overlap_text + "\n\n" + para
            else:
                # Paragraph itself is longer than max_chars; split it
                words = para.split()
                seg = ""
                for word in words:
                    if len(seg) + len(word) + 1 > max_chars:
                        chunks.append(seg.strip())
                        seg = word
                    else:
                        seg += (" " if seg else "") + word
                if seg:
                    current = seg
        else:
            current += ("\n\n" if current else "") + para
    if current:
        chunks.append(current.strip())
    return chunks


def extract_part_numbers(text: str) -> List[str]:
    """Extract likely Simplex part numbers from free text.

    Part numbers typically follow a pattern of 4 to 5 digits, a hyphen,
    and another 3 or 4 digits (e.g. ``4098-9714``).  This function
    returns all unique matches in the order encountered.
    """
    pattern = re.compile(r"\b\d{4,5}-\d{3,4}\b")
    matches = pattern.findall(text)
    seen = set()
    ordered = []
    for m in matches:
        if m not in seen:
            seen.add(m)
            ordered.append(m)
    return ordered


def parse_pdf(path: Path) -> Tuple[str, List[str]]:
    """
    Extract text and table data from a PDF.

    Attempts to use pdfplumber for table extraction and layout preservation.  If
    unavailable, falls back to PyMuPDF for basic text extraction.  Returns a
    tuple of (plain_text, tables) where tables is a list of strings
    representing rows.
    """
    plain_text = ""
    tables: List[str] = []
    # Try pdfplumber for structured extraction
    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    try:
                        plain_text += page.extract_text() or ""
                    except Exception:
                        # Some pages may not extract text
                        pass
                    # Extract tables if possible
                    try:
                        tbls = page.extract_tables()
                        for tbl in tbls:
                            # Join each row into a single string
                            for row in tbl:
                                if row:
                                    tables.append(" ".join([str(cell).strip() for cell in row if cell]))
                    except Exception:
                        pass
        except Exception as e:
            logger.warning(f"pdfplumber failed on {path}: {e}")
    # Fallback to PyMuPDF for text
    if not plain_text:
        if fitz is None:
            raise RuntimeError("Neither pdfplumber nor PyMuPDF installed; cannot parse PDF")
        try:
            doc = fitz.open(path)
            for page in doc:
                plain_text += page.get_text()
        except Exception as e:
            logger.error(f"PyMuPDF failed to parse {path}: {e}")
    return plain_text, tables


def parse_word(path: Path) -> Tuple[str, List[str]]:
    """
    Extract text and table data from a Word document.

    Uses python‑docx to extract paragraphs and tables.  Returns
    (plain_text, tables) where tables is a list of string representations
    of table rows.
    """
    if docx is None:
        raise RuntimeError("python-docx not installed; cannot parse Word")
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs]
    plain_text = "\n".join(paragraphs)
    tables: List[str] = []
    # Extract tables
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                tables.append(" ".join(cells))
    return plain_text, tables


def parse_excel(path: Path) -> Tuple[str, List[str]]:
    """
    Extract text and table data from an Excel workbook.

    Reads each sheet into a DataFrame and concatenates all cell values
    row by row.  Returns (plain_text, tables) where tables is a list of
    row strings to preserve some structure.
    """
    if pd is None:
        raise RuntimeError("pandas not installed; cannot parse Excel")
    book = pd.ExcelFile(path)
    plain_text_parts = []
    tables: List[str] = []
    for sheet_name in book.sheet_names:
        df = book.parse(sheet_name)
        df_str = df.astype(str).fillna("")
        for _, row in df_str.iterrows():
            row_text = " ".join(row.tolist())
            tables.append(row_text)
            plain_text_parts.append(row_text)
    return "\n".join(plain_text_parts), tables


def ingest_file(path: Path, llm: Optional[LLMInterface] = None) -> List[Component]:
    """Ingest a single datasheet and return a list of component definitions.

    The ingestion process attempts to extract part numbers and metadata
    from the file.  Text and tables are extracted using advanced
    parsers when available (pdfplumber, python‑docx tables, etc.).  The
    text is chunked to respect LLM context limits.  Each detected part
    number is associated with the smallest chunk containing it, so that
    only relevant text is sent to the language model.
    """
    ext = path.suffix.lower()
    try:
        if ext in [".pdf"]:
            plain_text, tables = parse_pdf(path)
        elif ext in [".docx", ".doc"]:
            plain_text, tables = parse_word(path)
        elif ext in [".xls", ".xlsx"]:
            plain_text, tables = parse_excel(path)
        else:
            logger.warning(f"Unsupported file type: {ext} for {path}")
            return []
    except Exception as e:
        logger.error(f"Failed to extract text from {path}: {e}")
        return []
    # Combine text and table rows into a single corpus for part number detection
    combined_text = plain_text + "\n" + "\n".join(tables)
    part_numbers = extract_part_numbers(combined_text)
    if not part_numbers:
        logger.warning(f"No part numbers found in {path}")
    # Chunk the plain text (not tables) for targeted extraction
    chunks = chunk_text(plain_text)
    # Map part numbers to chunks
    pn_to_chunk: Dict[str, str] = {}
    for pn in part_numbers:
        assigned = False
        for chunk in chunks:
            if pn in chunk:
                pn_to_chunk[pn] = chunk
                assigned = True
                break
        if not assigned:
            # fallback to first chunk or combined
            pn_to_chunk[pn] = chunks[0] if chunks else combined_text
    # Create component objects
    components: List[Component] = []
    for pn in part_numbers:
        comp = Component(
            part_number=pn,
            sku_type="product",
            category=ComponentType.ACCESSORY,
            description=f"Component {pn} extracted from {path.name}",
        )
        components.append(comp)
    # Use LLM to enrich if available
    if llm and llm.is_available() and part_numbers:
        tasks = []
        for comp in components:
            chunk = pn_to_chunk.get(comp.part_number, combined_text)
            prompt = f"""
You are processing a datasheet for a Simplex fire alarm component with part number {comp.part_number}.
Please extract the following information and return a JSON object with these keys:
  part_number, category, description, protocols, slot_size, compatible_with, requires, excludes,
  max_devices, certifications.
If any field is unknown, return null for that field.
Relevant excerpt of the datasheet:
""" + chunk
            tasks.append({"id": comp.part_number, "prompt": prompt})
        results = llm.batch_extract(tasks)
        for comp in components:
            result = results.get(comp.part_number)
            if result:
                try:
                    # Clean JSON from markdown code blocks if present
                    clean_result = result
                    if result.startswith('```json'):
                        clean_result = result.split('```json')[1].split('```')[0].strip()
                    elif result.startswith('```'):
                        clean_result = result.split('```')[1].split('```')[0].strip()
                    
                    data = json.loads(clean_result)
                    if data.get("category") in ComponentType.__members__:
                        comp.category = ComponentType[data["category"]]
                    comp.description = data.get("description", comp.description)
                    comp.protocols = [Protocol[p] for p in data.get("protocols", []) if p in Protocol.__members__]
                    comp.slot_size = data.get("slot_size", comp.slot_size)
                    comp.compatible_with = data.get("compatible_with", [])
                    comp.requires = data.get("requires", [])
                    comp.excludes = data.get("excludes", [])
                    if data.get("max_devices"):
                        comp.capacity_devices = data["max_devices"]
                    if data.get("certifications"):
                        comp.certifications = [CertificationType[c] for c in data.get("certifications", []) if c in CertificationType.__members__]
                except Exception as e:
                    logger.warning(f"Failed to parse LLM result for {comp.part_number}: {e}")
    return components


def ingest_datasheets(directory: Path, db: "SimplexDatabase", llm: Optional[LLMInterface] = None) -> None:
    """Ingest all datasheets in a directory into the database and graph."""
    if not directory.exists() or not directory.is_dir():
        logger.error(f"Invalid directory: {directory}")
        return
    for path in directory.iterdir():
        if path.suffix.lower() in [".pdf", ".doc", ".docx", ".xls", ".xlsx"]:
            comps = ingest_file(path, llm)
            for comp in comps:
                try:
                    comp.validate()
                    db.add_component(comp)
                except Exception as e:
                    logger.warning(f"Failed to add component {comp.part_number}: {e}")